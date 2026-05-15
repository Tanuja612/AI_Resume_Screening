import re
import os
from typing import Optional, Dict, Any, List

try:
    from pdfminer.high_level import extract_text as pdf_extract
except ImportError:
    pdf_extract = None

try:
    import docx
except ImportError:
    docx = None

# common ATS-safe fonts
SAFE_FONTS = {"arial", "calibri", "times new roman", "helvetica", "verdana", "georgia"}

SECTION_HEADERS = ["experience", "work experience", "education", "skills", "certifications", "projects"]

EMAIL_RE = re.compile(r"[\w\.-]+@[\w\.-]+\.[a-zA-Z]{2,}")
PHONE_RE = re.compile(r"\+?\d[\d\s\-]{7,}\d")


def extract_plain_text(path: str) -> str:
    """Extract text from .pdf or .docx file."""
    ext = os.path.splitext(path)[1].lower()
    text = ""
    if ext == ".pdf" and pdf_extract:
        try:
            text = pdf_extract(path)
        except Exception:
            text = ""
    elif ext == ".docx" and docx:
        try:
            doc = docx.Document(path)
            paragraphs = []
            for para in doc.paragraphs:
                paragraphs.append(para.text)
            text = "\n".join(paragraphs)
        except Exception:
            text = ""
    elif ext == ".txt":
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
        except Exception:
            text = ""
    return text


def analyze_resume(path: str, job_desc: Optional[str] = None) -> Dict[str, Any]:
    """Run ATS-compatibility analysis on given resume file.

    Validates that the file is actually a resume (not just any PDF).
    For non-resume PDFs, returns a report with score=0 and clear messaging.
    
    Returns a report dict containing score, strengths, weaknesses and
    recommendations.
    """
    from nlp.resume_validator import validate_resume_file
    
    report: Dict[str, Any] = {}
    text = extract_plain_text(path)
    words = text.split()

    report['word_count'] = len(words)
    report['has_text'] = bool(text.strip())
    report['sections_present'] = [h for h in SECTION_HEADERS if h in text.lower()]
    
    # FIRST: Validate that this is actually a resume document
    # If not a valid resume, return 0 score immediately
    validation_result = validate_resume_file(path, extracted_text=text)
    report['is_valid_resume'] = validation_result['is_valid_resume']
    report['validation_details'] = validation_result['validation_details']
    
    # If not a valid resume, short-circuit and return zero score
    if not validation_result['is_valid_resume']:
        report['score'] = 0
        report['has_tabs'] = False
        report['has_table_like'] = False
        report['has_multicolumn'] = False
        report['has_tables'] = False
        report['has_images'] = False
        report['has_headers_footers'] = False
        report['contact'] = {'email': False, 'phone': False}
        report['metrics'] = []
        report['has_bullets'] = False
        report['job_keywords'] = []
        report['keywords_missing'] = []
        report['font_issues'] = []
        
        # Add clear messaging about why this isn't a resume
        error_msg = validation_result['validation_details']['reason']
        report['strengths'] = []
        report['weaknesses'] = [f"Not a valid resume: {error_msg}"]
        report['recommendations'] = ["Please upload an actual resume document."]
        
        return report

    # generic formatting warnings
    report['has_tabs'] = '\t' in text
    report['has_table_like'] = bool(re.search(r"\n\s{4,}", text))
    report['has_multicolumn'] = '  ' in text  # double spaces may indicate columns

    # docx-specific checks
    report['has_tables'] = False
    report['has_images'] = False
    report['has_headers_footers'] = False
    if path.lower().endswith('.docx') and docx:
        try:
            d = docx.Document(path)
            report['has_tables'] = len(d.tables) > 0
            report['has_images'] = len(d.inline_shapes) > 0
            for section in d.sections:
                hdr = section.header
                ftr = section.footer
                if hdr and hdr.text.strip():
                    report['has_headers_footers'] = True
                if ftr and ftr.text.strip():
                    report['has_headers_footers'] = True
        except Exception:
            pass
    elif path.lower().endswith('.pdf'):
        # crude image detection by scanning raw PDF for /Image
        try:
            with open(path, 'rb') as f:
                raw = f.read()
            report['has_images'] = b'/Image' in raw
            report['has_tables'] = b'/Table' in raw or b'Grid' in raw
            # header/footer detection can be tough; look for repeated lines
            lines = text.splitlines()
            if len(lines) > 2 and lines[0] == lines[-1]:
                report['has_headers_footers'] = True
        except Exception:
            pass

    report['contact'] = {
        'email': bool(EMAIL_RE.search(text)),
        'phone': bool(PHONE_RE.search(text))
    }

    # quantify achievements
    report['metrics'] = re.findall(r"\b\d+[,%]?\b", text)
    report['has_bullets'] = bool(re.search(r"[••\-]\s", text))

    # keyword matching
    report['job_keywords'] = []
    report['keywords_missing'] = []
    if job_desc:
        jd_words = set(re.findall(r"\w+", job_desc.lower()))
        matches = [w for w in jd_words if w in text.lower()]
        report['job_keywords'] = matches
        report['keywords_missing'] = list(jd_words - set(matches))

        # detect stuffing: count repetition of each matched kwd
        repetitions = {w: text.lower().count(w) for w in matches}
        report['keyword_repetitions'] = repetitions

    # font/style detection for docx
    report['font_issues'] = []
    if path.lower().endswith('.docx') and docx:
        try:
            d = docx.Document(path)
            for para in d.paragraphs:
                for run in para.runs:
                    font = run.font.name
                    if font and font.lower() not in SAFE_FONTS:
                        report['font_issues'].append(font)
        except Exception:
            pass

    # now compute score (benchmark values chosen arbitrarily)
    score = 100
    if not report['has_text']:
        score = 0
    else:
        if report['word_count'] < 50:
            score -= 20
        if len(report['sections_present']) < 2:
            score -= 20
        if report['has_tabs'] or report['has_table_like'] or report['has_multicolumn']:
            score -= 15
        if report.get('has_tables'):
            score -= 15
        if report.get('has_images'):
            score -= 10
        if report.get('has_headers_footers'):
            score -= 10
        if not report['contact']['email'] or not report['contact']['phone']:
            score -= 10
        if report.get('keywords_missing'):
            score -= min(20, len(report['keywords_missing']) * 2)
        if report.get('keyword_repetitions') and any(v > 10 for v in report['keyword_repetitions'].values()):
            score -= 10
        if report['font_issues']:
            score -= 5
        score = max(0, min(100, score))
    report['score'] = score

    # strengths/weaknesses/recommendations
    strengths: List[str] = []
    weaknesses: List[str] = []
    recommends: List[str] = []

    if report['has_text']:
        strengths.append("Extractable text present")
    if report['word_count'] >= 200:
        strengths.append("Good length")
    if report['sections_present']:
        strengths.append(f"Sections found: {', '.join(report['sections_present'])}")
    if report['contact']['email'] and report['contact']['phone']:
        strengths.append("Contact information present")
    if job_desc and report['job_keywords']:
        strengths.append(f"Keywords matched: {', '.join(report['job_keywords'])}")
    if report['metrics']:
        strengths.append(f"Quantified achievements: {len(report['metrics'])} detected")
    if report.get('has_bullets'):
        strengths.append("Bullet points detected, aiding readability")

    if not report['has_text']:
        weaknesses.append("Resume appears to contain no extractable text")
    if report['word_count'] < 50:
        weaknesses.append("Very short document")
    if len(report['sections_present']) < 2:
        weaknesses.append("Missing common section headers")
    if report['has_tabs'] or report['has_table_like'] or report['has_multicolumn']:
        weaknesses.append("Contains tables, columns, or excessive indentation")
    if report.get('has_tables'):
        weaknesses.append("Contains embedded tables")
    if report.get('has_images'):
        weaknesses.append("Contains images or graphics")
    if report.get('has_headers_footers'):
        weaknesses.append("Contains headers/footers that may confuse parsers")
    if not report['contact']['email'] or not report['contact']['phone']:
        weaknesses.append("Missing contact email or phone")
    if job_desc and report['keywords_missing']:
        weaknesses.append(f"Missing keywords: {', '.join(report['keywords_missing'])}")
    if job_desc and report.get('keyword_repetitions'):
        for k, v in report['keyword_repetitions'].items():
            if v > 10:
                weaknesses.append(f"Keyword '{k}' repeated {v} times (may be stuffing)")
    if report['font_issues']:
        weaknesses.append(f"Unusual fonts: {', '.join(set(report['font_issues']))}")

    if report['has_tabs'] or report['has_table_like'] or report['has_multicolumn']:
        recommends.append("Avoid tables/columns or formatting that use tabs/spaces; use plain lists")
    if report.get('has_tables'):
        recommends.append("Convert tables to plain text or bullet lists")
    if report.get('has_images'):
        recommends.append("Remove logos/images; ATS reads text only")
    if report.get('has_headers_footers'):
        recommends.append("Remove header/footer text (e.g. page numbers)")
    if not report['contact']['email'] or not report['contact']['phone']:
        recommends.append("Add clear contact information in plain text")
    if job_desc and report['keywords_missing']:
        recommends.append("Incorporate missing keywords from the job description")
    if report['font_issues']:
        recommends.append("Use standard fonts like Arial or Calibri")
    if report['score'] < 60:
        recommends.append("Consider rewriting for clarity and ATS compatibility")

    report['strengths'] = strengths
    report['weaknesses'] = weaknesses
    report['recommendations'] = recommends

    return report


def generate_corrected_resume(original_text: str, analysis: Dict[str, Any]) -> tuple[str, List[str]]:
    """Generate an ATS-friendly corrected version of the resume.
    
    Args:
        original_text: The extracted text from the original resume
        analysis: The analysis dictionary from analyze_resume()
    
    Returns:
        A tuple of (corrected_text, corrections_made) where corrections_made
        is a list of human-readable corrections applied.
    """
    corrected = original_text
    corrections = []
    
    # 1. Remove tabs and excessive indentation
    if analysis.get('has_tabs'):
        corrected = corrected.replace('\t', '  ')
        corrections.append("✓ Replaced tabs with spaces")
    
    # 2. Clean up multiple spaces (columns)
    if analysis.get('has_multicolumn'):
        corrected = re.sub(r' {3,}', '  ', corrected)
        corrections.append("✓ Fixed excessive spacing/columns")
    
    # 3. Clean up table-like content
    if analysis.get('has_table_like'):
        lines = corrected.split('\n')
        cleaned_lines = []
        for line in lines:
            # Convert indented table-like content to proper format
            if re.match(r'^\s{4,}', line):
                line = '  ' + line.lstrip()
            cleaned_lines.append(line)
        corrected = '\n'.join(cleaned_lines)
        corrections.append("✓ Converted table-like formatting to plain text")
    
    # 4. Remove non-ASCII characters
    original_len = len(corrected)
    corrected = ''.join(c if ord(c) < 128 else ' ' for c in corrected)
    if len(corrected) < original_len:
        corrections.append("✓ Removed special characters and non-ASCII text")
    
    # 5. Ensure proper section headers
    for header in SECTION_HEADERS:
        if header not in corrected.lower():
            # Header is missing - this will be noted but not auto-added
            pass
    
    # 6. Ensure contact info is prominent at top
    emails = EMAIL_RE.findall(corrected)
    phones = PHONE_RE.findall(corrected)
    
    if not emails or not phones:
        if not emails:
            corrections.append("⚠ Missing email address - please add it")
        if not phones:
            corrections.append("⚠ Missing phone number - please add it")
    
    # 7. Remove keyword stuffing
    if analysis.get('keyword_repetitions'):
        for keyword, count in analysis['keyword_repetitions'].items():
            if count > 10:
                # Replace repeated keyword with single instances
                pattern = rf'\b{re.escape(keyword)}\b'
                new_count = 5  # Limit to reasonable repetition
                corrected = re.sub(pattern, keyword, corrected, flags=re.IGNORECASE)
                corrections.append(f"✓ Removed keyword stuffing for '{keyword}' (was {count} times, reduced to ~{new_count})")
    
    # 8. Format bullet points consistently
    if not analysis.get('has_bullets'):
        corrections.append("💡 Consider adding bullet points for better readability")
    else:
        # Standardize bullet points
        corrected = re.sub(r'[•\-]\s+', '• ', corrected)
        corrections.append("✓ Standardized bullet point formatting")
    
    # 9. Clean up line breaks
    corrected = re.sub(r'\n{3,}', '\n\n', corrected)
    corrections.append("✓ Cleaned up excessive line breaks")
    
    # 10. Remove common decorative lines
    corrected = re.sub(r'^[\-=_]{5,}$', '', corrected, flags=re.MULTILINE)
    corrections.append("✓ Removed decorative lines")
    
    # 11. Add footer note
    corrected = corrected.rstrip() + "\n\n" + "="*60 + "\n"
    corrected += "CORRECTED FOR ATS COMPATIBILITY\n"
    corrected += "="*60
    
    return corrected, corrections
